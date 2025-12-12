import docx
import os

def inspect_structure(filename):
    print(f"Current WD: {os.getcwd()}")
    if not os.path.exists(filename):
        print(f"File not found: {filename}")
        return

    try:
        doc = docx.Document(filename)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # Print headers or potential section markers
            if len(text) < 20 and any(k in text for k in ["单选题", "多选题", "判断题", "简答题", "填空题", "一、", "二、", "三、", "四、"]):
                print(f"[{i}] {text}")
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    inspect_structure("毛概（2025-2026-1）期末练习题库.docx")